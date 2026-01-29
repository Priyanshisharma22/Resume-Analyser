import os
import re
from docx import Document
from fpdf import FPDF
from pypdf import PdfReader

OUTPUT_DIR = "outputs"
os.makedirs(OUTPUT_DIR, exist_ok=True)

def read_docx(file) -> str:
    doc = Document(file)
    return "\n".join([p.text for p in doc.paragraphs]).strip()

def read_pdf(file) -> str:
    reader = PdfReader(file)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text() or ""
        text += page_text + "\n"
    return text.strip()

def save_docx(text: str, filename: str) -> str:
    path = os.path.join(OUTPUT_DIR, filename)
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(path)
    return path

def _wrap_long_tokens(text: str, max_len: int = 60) -> str:
    def break_token(token):
        if len(token) <= max_len:
            return token
        return "\n".join(token[i:i + max_len] for i in range(0, len(token), max_len))
    parts = text.split()
    parts = [break_token(p) for p in parts]
    return " ".join(parts)

def save_pdf(text: str, filename: str) -> str:
    path = os.path.join(OUTPUT_DIR, filename)

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=12)
    pdf.add_page()
    pdf.set_font("Helvetica", size=11)

    text = text.replace("\t", "    ")
    text = re.sub(r"[^\x00-\x7F]+", " ", text)
    text = _wrap_long_tokens(text, max_len=60)

    for line in text.split("\n"):
        if not line.strip():
            pdf.ln(5)
        else:
            pdf.multi_cell(0, 6, line)

    pdf.output(path)
    return path
